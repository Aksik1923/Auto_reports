import os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX  # Импортируем предопределённые цвета
from docx.shared import RGBColor  # Импортируем RGBColor

def find_document(base_path, subfolder, document_name):
    # Поиск документа в указанной подпапке
    target_folder = os.path.join(base_path, subfolder)
    if os.path.exists(target_folder):
        for root, dirs, files in os.walk(target_folder):
            if document_name in files:
                return os.path.join(root, document_name)
    return None

def edit_introduction(doc_path, tz_path, customer_path, executor_path, contract_path, customers_base_path):
    try:
        # Открываем документ
        doc = Document(doc_path)
        
        # Находим начало и конец главы "Введение"
        start_intro = None
        end_intro = None
        for i, paragraph in enumerate(doc.paragraphs):
            if "ВВЕДЕНИЕ" in paragraph.text and paragraph.style.name == "01 Заголовок":
                start_intro = i
            if "ФИЗИКО-ГЕОГРАФИЧЕСКАЯ ХАРАКТЕРИСТИКА РАЙОНА РАБОТ" in paragraph.text and paragraph.style.name == "01 Заголовок":
                end_intro = i
                break
        
        if start_intro is None or end_intro is None:
            print("Ошибка: Не удалось найти границы главы 'Введение'.")
            return
        
        # Редактируем главу "Введение"
        for i in range(start_intro, end_intro):
            paragraph = doc.paragraphs[i]
            
            # 1. Местоположение объекта
            if "Местоположение объекта:" in paragraph.text:
                location = get_table_value(tz_path, "Месторасположение")
                if location:
                    replace_text_preserve_formatting(paragraph, "Местоположение объекта:", f"Местоположение объекта: {location}")
            
            # 2. Название объекта
            if "Название объекта:" in paragraph.text:
                object_name = get_table_value(tz_path, "Наименование объекта")
                if object_name:
                    replace_text_preserve_formatting(paragraph, "Название объекта:", f"Название объекта: {object_name}")
            
            # 3. Вид строительства
            if "Вид строительства:" in paragraph.text:
                construction_type = get_table_value(tz_path, "Вид строительства")
                if construction_type:
                    replace_text_preserve_formatting(paragraph, "Вид строительства:", f"Вид строительства: {construction_type}")
                else:
                    # Удаляем строку, если вид строительства не найден
                    paragraph.text = paragraph.text.replace("Вид строительства:", "")
            
            # 4. Стадия проектирования
            if "Стадия проектирования:" in paragraph.text:
                design_stage = get_table_value(tz_path, "Стадия проектирования")
                if not design_stage:
                    design_stage = search_text_in_tz(tz_path, ["проектная документация", "предпроектная документация"])
                if design_stage:
                    replace_text_preserve_formatting(paragraph, "Стадия проектирования:", f"Стадия проектирования: {design_stage}")
            
            # 5. Заказчик
            if "Заказчик" in paragraph.text and paragraph.runs[0].bold:
                customer_name = get_customer_name(customer_path)
                if customer_name:
                    customer_info = get_customer_info(customer_name, customers_base_path)
                    if customer_info:
                        replace_text_preserve_formatting(paragraph, "Заказчик", f"Заказчик: {customer_info}")
            
            # 6. Изыскательская организация
            if "Изыскательская организация" in paragraph.text and paragraph.runs[0].bold:
                organization = choose_organization(customers_base_path)
                if organization:
                    replace_text_preserve_formatting(paragraph, "Изыскательская организация", f"Изыскательская организация: {organization}")
            
            # 7. Сроки проведения работ
            if "Сроки проведения работ:" in paragraph.text:
                deadlines = get_table_value(tz_path, "Сроки выполнения работ")
                if deadlines:
                    replace_text_preserve_formatting(paragraph, "Сроки проведения работ:", f"Сроки проведения работ: {deadlines}")
            
            # 8. Основные объекты проектирования
            if "Основные объекты проектирования:" in paragraph.text:
                objects = get_lines_with_plus(tz_path, "Перечень объектов проектирования")
                if objects:
                    # Форматируем объекты как маркированный список
                    formatted_objects = format_as_bullet_list(objects)
                    replace_text_preserve_formatting(paragraph, "Основные объекты проектирования:", f"Основные объекты проектирования:\n{formatted_objects}")
            
            # 9. Система координат
            if "Система координат:" in paragraph.text:
                coordinate_system = get_table_value(tz_path, "Система координат")
                if coordinate_system:
                    replace_text_preserve_formatting(paragraph, "Система координат:", f"Система координат: {coordinate_system}")
                else:
                    replace_text_preserve_formatting(paragraph, "Система координат:", "Система координат: системы координат в текущем ТЗ нет")
                    highlight_text(paragraph, "системы координат в текущем ТЗ нет", WD_COLOR_INDEX.YELLOW)  # Жёлтый цвет
            
            # 10. Система высот
            if "Система высот" in paragraph.text and paragraph.runs[0].bold:
                height_system = get_table_value(tz_path, "Система высот")
                if height_system:
                    replace_text_preserve_formatting(paragraph, "Система высот", f"Система высот: {height_system}")
                else:
                    replace_text_preserve_formatting(paragraph, "Система высот", "Система высот: системы высот в текущем ТЗ нет")
                    highlight_text(paragraph, "системы высот в текущем ТЗ нет", WD_COLOR_INDEX.YELLOW)  # Жёлтый цвет
            
            # 11. Наименование объекта (вчерашняя процедура)
            if "по объекту:" in paragraph.text:
                object_name = get_table_value(tz_path, "Наименование объекта")
                if object_name:
                    # Убираем точку в конце, если она есть
                    if object_name.endswith("."):
                        object_name = object_name[:-1]
                    replace_text_without_bold(paragraph, "по объекту:", f"по объекту: {object_name}")
            
            # 12. Заказчик и генеральный директор (вчерашняя процедура)
            if "генеральным директором" in paragraph.text:
                customer_name = get_customer_name(customer_path)
                if customer_name:
                    director_name = get_director_name(customer_name, customers_base_path)
                    replace_text_without_bold(paragraph, "генеральным директором", f"генеральным директором {customer_name} {director_name}")
            
            # 13. Исполнитель (вчерашняя процедура)
            if "Договор между" in paragraph.text:
                customer_name = get_customer_name(customer_path)
                executor_name = get_executor_name(executor_path)
                if customer_name and executor_name:
                    replace_text_without_bold(paragraph, "Договор между", f"Договор между {customer_name} и {executor_name}")
            
            # 14. Договор (вчерашняя процедура)
            if "и выданное Техническое задание" in paragraph.text:
                contract_text = get_contract_text(contract_path)
                if contract_text:
                    replace_text_without_bold(paragraph, "и выданное Техническое задание", f"{contract_text} и выданное Техническое задание")
        
        # Создаём папку "Готовый отчёт", если её нет
        output_folder = os.path.join(os.path.dirname(doc_path), "Готовый отчёт")
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
        
        # Сохраняем отредактированный документ в папку "Готовый отчёт"
        output_path = os.path.join(output_folder, os.path.basename(doc_path))
        doc.save(output_path)
        print(f"Глава 'Введение' успешно отредактирована. Отчёт сохранён в: {output_path}")
    except Exception as e:
        print(f"Ошибка при редактировании документа: {e}")

def get_table_value(tz_path, keyword):
    try:
        tz_doc = Document(tz_path)
        for table in tz_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if keyword in cell.text:
                        return row.cells[-1].text.strip()  # Берём значение из последнего столбца
        return ""
    except Exception as e:
        print(f"Ошибка при чтении таблицы в ТЗ: {e}")
        return ""

def get_lines_with_plus(tz_path, keyword):
    """
    Извлекает из таблицы ТЗ строки, содержащие знак "+".
    :param tz_path: Путь к документу ТЗ.
    :param keyword: Ключевое слово для поиска ячейки.
    :return: Список строк, содержащих знак "+".
    """
    try:
        tz_doc = Document(tz_path)
        for table in tz_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if keyword in cell.text:
                        # Извлекаем текст из ячейки
                        text = cell.text.strip()
                        # Разделяем текст на строки
                        lines = text.split("\n")
                        # Фильтруем строки, содержащие знак "+"
                        plus_lines = [line.strip() for line in lines if "+" in line]
                        return plus_lines
        return []
    except Exception as e:
        print(f"Ошибка при чтении таблицы в ТЗ: {e}")
        return []

def format_as_bullet_list(lines):
    """
    Форматирует список строк как маркированный список.
    :param lines: Список строк.
    :return: Отформатированный текст с маркерами.
    """
    # Добавляем маркер к каждой строке
    formatted_lines = [f"• {line}" for line in lines if line.strip()]
    # Объединяем строки с переносами
    return "\n".join(formatted_lines)

def search_text_in_tz(tz_path, keywords):
    try:
        tz_doc = Document(tz_path)
        for paragraph in tz_doc.paragraphs:
            for keyword in keywords:
                if keyword in paragraph.text:
                    return keyword
        return ""
    except Exception as e:
        print(f"Ошибка при поиске текста в ТЗ: {e}")
        return ""

def get_customer_name(customer_path):
    try:
        with open(customer_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except Exception as e:
        print(f"Ошибка при чтении файла Заказчик.txt: {e}")
        return ""

def get_customer_info(customer_name, customers_base_path):
    try:
        customer_dir = os.path.join(customers_base_path, "Заказчики", customer_name)
        info_path = os.path.join(customer_dir, "Информация о заказчике.txt")
        if os.path.exists(info_path):
            with open(info_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        else:
            print(f"Ошибка: Файл 'Информация о заказчике.txt' не найден в папке '{customer_dir}'.")
            return ""
    except Exception as e:
        print(f"Ошибка при чтении информации о заказчике: {e}")
        return ""

def get_director_name(customer_name, customers_base_path):
    try:
        customer_dir = os.path.join(customers_base_path, "Заказчики", customer_name)
        director_path = os.path.join(customer_dir, "Генеральный директор.txt")
        if os.path.exists(director_path):
            with open(director_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        else:
            print(f"Ошибка: Файл 'Генеральный директор.txt' не найден в папке '{customer_dir}'.")
            return ""
    except Exception as e:
        print(f"Ошибка при чтении информации о генеральном директоре: {e}")
        return ""

def get_executor_name(executor_path):
    try:
        with open(executor_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except Exception as e:
        print(f"Ошибка при чтении файла Исполнитель.txt: {e}")
        return ""

def get_contract_text(contract_path):
    try:
        with open(contract_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except Exception as e:
        print(f"Ошибка при чтении файла Договор.txt: {e}")
        return ""

def choose_organization(customers_base_path):
    try:
        organizations = ["МСЛ", "МОРИИ"]
        print("Выберите изыскательскую организацию:")
        for i, org in enumerate(organizations):
            print(f"{i + 1}. {org}")
        choice = int(input("Введите номер: ")) - 1
        if 0 <= choice < len(organizations):
            org_path = os.path.join(customers_base_path, "Исполнители", organizations[choice], "Адрес.txt")
            if os.path.exists(org_path):
                with open(org_path, 'r', encoding='utf-8') as file:
                    return file.read().strip()
            else:
                print(f"Ошибка: Файл 'Адрес.txt' не найден в папке '{org_path}'.")
                return ""
        else:
            print("Ошибка: Неверный выбор.")
            return ""
    except Exception as e:
        print(f"Ошибка при выборе организации: {e}")
        return ""

def highlight_text(paragraph, text, color):
    """
    Подсвечивает текст в параграфе указанным цветом.
    :param paragraph: Параграф, в котором нужно подсветить текст.
    :param text: Текст, который нужно подсветить.
    :param color: Цвет из WD_COLOR_INDEX (например, WD_COLOR_INDEX.YELLOW).
    """
    for run in paragraph.runs:
        if text in run.text:
            run.font.highlight_color = color

def replace_text_preserve_formatting(paragraph, old_text, new_text):
    """
    Заменяет текст в параграфе, сохраняя форматирование ключевых слов.
    :param paragraph: Параграф, в котором нужно заменить текст.
    :param old_text: Текст, который нужно заменить.
    :param new_text: Новый текст.
    """
    if old_text in paragraph.text:
        # Разделяем текст на части до и после замены
        parts = paragraph.text.split(old_text)
        if len(parts) != 2:
            return
        
        # Очищаем параграф
        paragraph.clear()
        
        # Добавляем первую часть (полужирный стиль)
        run = paragraph.add_run(parts[0])
        run.bold = True
        
        # Добавляем старый текст (полужирный стиль)
        run = paragraph.add_run(old_text)
        run.bold = True
        
        # Добавляем новый текст (обычный стиль)
        run = paragraph.add_run(new_text.replace(old_text, ""))
        run.bold = False
        
        # Добавляем вторую часть (полужирный стиль)
        run = paragraph.add_run(parts[1])
        run.bold = True

def replace_text_without_bold(paragraph, old_text, new_text):
    """
    Заменяет текст в параграфе, не сохраняя полужирный стиль для вставленного текста.
    :param paragraph: Параграф, в котором нужно заменить текст.
    :param old_text: Текст, который нужно заменить.
    :param new_text: Новый текст.
    """
    if old_text in paragraph.text:
        # Разделяем текст на части до и после замены
        parts = paragraph.text.split(old_text)
        if len(parts) != 2:
            return
        
        # Очищаем параграф
        paragraph.clear()
        
        # Добавляем первую часть (обычный стиль)
        run = paragraph.add_run(parts[0])
        run.bold = False
        
        # Добавляем старый текст (обычный стиль)
        run = paragraph.add_run(old_text)
        run.bold = False
        
        # Добавляем новый текст (обычный стиль)
        run = paragraph.add_run(new_text.replace(old_text, ""))
        run.bold = False
        
        # Добавляем вторую часть (обычный стиль)
        run = paragraph.add_run(parts[1])
        run.bold = False

if __name__ == "__main__":
    base_path = input("Введите путь к базовой папке: ")  # Путь к базовой папке
    customers_base_path = input("Введите путь к папке с заказчиками: ")  # Путь к папке с заказчиками
    subfolder = "Пояснительная записка"  # Подпапка, где ищем документ "ПО ИГФИ"
    document_name = "ПО ИГФИ.docx"
    doc_path = find_document(base_path, subfolder, document_name)
    
    if doc_path:
        tz_path = os.path.join(base_path, "ТЗ", "ТЗ.docx")
        customer_path = os.path.join(base_path, "Заказчик", "Заказчик.txt")
        executor_path = os.path.join(base_path, "Исполнитель", "Исполнитель.txt")
        contract_path = os.path.join(base_path, "Договор", "Договор.txt")
        
        # Проверка существования всех необходимых файлов
        if not os.path.exists(tz_path):
            print(f"Ошибка: Файл ТЗ не найден по пути '{tz_path}'.")
        elif not os.path.exists(customer_path):
            print(f"Ошибка: Файл Заказчик.txt не найден по пути '{customer_path}'.")
        elif not os.path.exists(executor_path):
            print(f"Ошибка: Файл Исполнитель.txt не найден по пути '{executor_path}'.")
        elif not os.path.exists(contract_path):
            print(f"Ошибка: Файл Договор.txt не найден по пути '{contract_path}'.")
        else:
            edit_introduction(doc_path, tz_path, customer_path, executor_path, contract_path, customers_base_path)
    else:
        print(f"Документ '{document_name}' не найден в папке '{subfolder}'.")